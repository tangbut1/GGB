from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import json


class DOCXReportGenerator:
    """DOCX报告生成器 - 生成可编辑的分析报告"""
    
    def __init__(self):
        self.doc = None
    
    def create_report(self, sentiment_summary: Dict[str, Any],
                     trend_summary: Dict[str, Any],
                     analyzed_news: List[Dict[str, Any]],
                     trend_data: Dict[str, Any],
                     output_path: str = "results/reports/market_analysis_report.docx") -> str:
        """
        创建完整的DOCX报告
        
        Args:
            sentiment_summary: 情绪分析摘要
            trend_summary: 趋势分析摘要
            analyzed_news: 已分析的新闻数据
            trend_data: 趋势预测数据
            output_path: 输出路径
            
        Returns:
            生成的DOCX文件路径
        """
        # 创建文档
        self.doc = Document()
        
        # 设置文档样式
        self._setup_document_style()
        
        # 添加标题页
        self._add_title_page(sentiment_summary)
        
        # 添加目录
        self._add_table_of_contents()
        
        # 添加执行摘要
        self._add_executive_summary(sentiment_summary, trend_summary)
        
        # 添加情绪分析详情
        self._add_sentiment_analysis(sentiment_summary, analyzed_news)
        
        # 添加趋势预测
        self._add_trend_prediction(trend_summary, trend_data)
        
        # 添加新闻详情
        self._add_news_details(analyzed_news)
        
        # 添加结论和建议
        self._add_conclusions(sentiment_summary, trend_summary)
        
        # 保存文档
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(output_path)
        
        print(f"✅ DOCX报告已生成: {output_path}")
        return output_path
    
    def _setup_document_style(self):
        """设置文档样式"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(12)
    
    def _add_title_page(self, sentiment_summary: Dict[str, Any]):
        """添加标题页"""
        # 主标题
        title = self.doc.add_heading('MarketPulse 智能市场分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 副标题
        subtitle = self.doc.add_heading('基于AI的市场情绪分析与趋势预测', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 生成时间
        current_time = datetime.now().strftime("%Y年%m月%d日 %H:%M")
        time_para = self.doc.add_paragraph(f"报告生成时间: {current_time}")
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加空行
        self.doc.add_paragraph()
        
        # 关键指标表格
        self.doc.add_heading('关键指标概览', level=2)
        
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标'
        hdr_cells[1].text = '数值'
        
        # 添加数据行
        metrics = [
            ("总新闻数", str(sentiment_summary.get('total_news', 0))),
            ("积极新闻", str(sentiment_summary.get('positive_count', 0))),
            ("消极新闻", str(sentiment_summary.get('negative_count', 0))),
            ("平均情绪得分", f"{sentiment_summary.get('avg_sentiment', 0):.3f}")
        ]
        
        for metric, value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
    
    def _add_table_of_contents(self):
        """添加目录"""
        self.doc.add_heading('目录', level=1)
        
        toc_items = [
            "1. 执行摘要",
            "2. 情绪分析详情", 
            "3. 趋势预测分析",
            "4. 新闻详情分析",
            "5. 结论与建议"
        ]
        
        for item in toc_items:
            self.doc.add_paragraph(item, style='List Number')
    
    def _add_executive_summary(self, sentiment_summary: Dict[str, Any], 
                              trend_summary: Dict[str, Any]):
        """添加执行摘要"""
        self.doc.add_heading('1. 执行摘要', level=1)
        
        # 情绪分析摘要
        total_news = sentiment_summary.get('total_news', 0)
        positive_count = sentiment_summary.get('positive_count', 0)
        negative_count = sentiment_summary.get('negative_count', 0)
        avg_sentiment = sentiment_summary.get('avg_sentiment', 0)
        
        summary_text = f"""
本报告基于对 {total_news} 条财经新闻的深度分析，通过多模型融合的情绪分析技术，
识别出积极新闻 {positive_count} 条，消极新闻 {negative_count} 条，
整体市场情绪得分为 {avg_sentiment:.3f}。
        """
        
        self.doc.add_paragraph(summary_text.strip())
        
        # 趋势分析摘要
        if trend_summary.get('status') == 'success':
            trend_direction = trend_summary.get('trend_direction', 'neutral')
            confidence = trend_summary.get('confidence', 0)
            
            trend_text = f"""
基于Prophet时间序列模型的趋势预测显示，市场情绪呈{trend_direction}趋势，
预测置信度为 {confidence:.1%}。
            """
            
            self.doc.add_paragraph(trend_text.strip())
    
    def _add_sentiment_analysis(self, sentiment_summary: Dict[str, Any], 
                               analyzed_news: List[Dict[str, Any]]):
        """添加情绪分析详情"""
        self.doc.add_heading('2. 情绪分析详情', level=1)
        
        # 情绪分布统计
        dist = sentiment_summary.get('sentiment_distribution', {})
        
        self.doc.add_heading('2.1 情绪分布统计', level=2)
        
        # 创建情绪分布表格
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # 表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '情绪类型'
        hdr_cells[1].text = '数量'
        hdr_cells[2].text = '占比'
        
        # 添加数据
        total_news = sentiment_summary.get('total_news', 1)
        for emotion, count in dist.items():
            row_cells = table.add_row().cells
            row_cells[0].text = emotion
            row_cells[1].text = str(count)
            row_cells[2].text = f"{count / total_news * 100:.1f}%"
        
        # 情绪分析
        self.doc.add_heading('2.2 情绪分析', level=2)
        
        avg_sentiment = sentiment_summary.get('avg_sentiment', 0)
        if avg_sentiment > 0.1:
            sentiment_analysis = "市场整体情绪偏向积极，投资者信心较强。"
        elif avg_sentiment < -0.1:
            sentiment_analysis = "市场整体情绪偏向消极，投资者信心不足。"
        else:
            sentiment_analysis = "市场整体情绪相对中性，投资者保持观望态度。"
        
        self.doc.add_paragraph(f"情绪分析: {sentiment_analysis}")
    
    def _add_trend_prediction(self, trend_summary: Dict[str, Any], 
                            trend_data: Dict[str, Any]):
        """添加趋势预测分析"""
        self.doc.add_heading('3. 趋势预测分析', level=1)
        
        if trend_summary.get('status') == 'error':
            error_text = f"趋势预测失败: {trend_summary.get('message', '未知错误')}"
            self.doc.add_paragraph(error_text)
            return
        
        # 预测结果
        trend_direction = trend_summary.get('trend_direction', 'neutral')
        confidence = trend_summary.get('confidence', 0)
        data_points = trend_summary.get('data_points', 0)
        forecast_periods = trend_summary.get('forecast_periods', 0)
        
        self.doc.add_heading('3.1 预测结果', level=2)
        
        # 创建预测结果表格
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标'
        hdr_cells[1].text = '数值'
        
        prediction_data = [
            ("数据点数", str(data_points)),
            ("趋势方向", trend_direction),
            ("预测置信度", f"{confidence:.1%}"),
            ("预测天数", str(forecast_periods))
        ]
        
        for metric, value in prediction_data:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
        
        # 投资建议
        self.doc.add_heading('3.2 投资建议', level=2)
        recommendation = trend_summary.get('recommendation', '')
        if recommendation:
            self.doc.add_paragraph(f"投资建议: {recommendation}")
    
    def _add_news_details(self, analyzed_news: List[Dict[str, Any]]):
        """添加新闻详情"""
        self.doc.add_heading('4. 新闻详情分析', level=1)
        
        # 显示前10条重要新闻
        important_news = sorted(analyzed_news, 
                              key=lambda x: abs(x.get('sentiment_score', 0)), 
                              reverse=True)[:10]
        
        self.doc.add_heading('4.1 重要新闻分析', level=2)
        
        for i, news in enumerate(important_news, 1):
            title = news.get('title', '无标题')
            sentiment_score = news.get('sentiment_score', 0)
            sentiment_label = news.get('sentiment_label', 'neutral')
            confidence = news.get('sentiment_confidence', 0)
            publish_time = news.get('publish_time', '未知')
            source = news.get('source', '未知')
            
            # 新闻标题
            self.doc.add_heading(f'新闻 {i}: {title}', level=3)
            
            # 新闻信息表格
            table = self.doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'
            
            info_data = [
                ("情绪得分", f"{sentiment_score:.3f}"),
                ("情绪标签", sentiment_label),
                ("置信度", f"{confidence:.3f}"),
                ("发布时间", publish_time)
            ]
            
            for j, (label, value) in enumerate(info_data):
                row_cells = table.rows[j].cells
                row_cells[0].text = label
                row_cells[1].text = value
            
            # 新闻内容摘要
            content = news.get('content', '无内容')
            if content:
                self.doc.add_paragraph(f"内容摘要: {content[:200]}...")
            
            self.doc.add_paragraph()  # 添加空行
    
    def _add_conclusions(self, sentiment_summary: Dict[str, Any], 
                        trend_summary: Dict[str, Any]):
        """添加结论与建议"""
        self.doc.add_heading('5. 结论与建议', level=1)
        
        # 综合结论
        avg_sentiment = sentiment_summary.get('avg_sentiment', 0)
        total_news = sentiment_summary.get('total_news', 0)
        
        conclusion_text = f"""
基于对 {total_news} 条新闻的深度分析，当前市场情绪得分为 {avg_sentiment:.3f}。
        """
        
        self.doc.add_paragraph(conclusion_text.strip())
        
        # 投资建议
        if trend_summary.get('status') == 'success':
            trend_direction = trend_summary.get('trend_direction', 'neutral')
            
            if trend_direction == 'positive':
                advice = "建议关注市场机会，适当增加投资配置。"
            elif trend_direction == 'negative':
                advice = "建议谨慎投资，适当降低风险敞口。"
            else:
                advice = "建议保持观望，等待更明确的市场信号。"
            
            self.doc.add_paragraph(f"投资建议: {advice}")
        
        # 风险提示
        self.doc.add_heading('风险提示', level=2)
        risk_warning = """
本报告基于历史数据和模型预测，仅供参考，不构成投资建议。
投资有风险，决策需谨慎。
        """
        
        self.doc.add_paragraph(risk_warning.strip())


def generate_docx_report(sentiment_summary: Dict[str, Any],
                       trend_summary: Dict[str, Any],
                       analyzed_news: List[Dict[str, Any]],
                       trend_data: Dict[str, Any],
                       output_path: str = "results/reports/market_analysis_report.docx") -> str:
    """
    便捷函数：生成DOCX报告
    
    Args:
        sentiment_summary: 情绪分析摘要
        trend_summary: 趋势分析摘要
        analyzed_news: 已分析的新闻数据
        trend_data: 趋势预测数据
        output_path: 输出路径
        
    Returns:
        生成的DOCX文件路径
    """
    generator = DOCXReportGenerator()
    return generator.create_report(sentiment_summary, trend_summary, 
                                 analyzed_news, trend_data, output_path)
